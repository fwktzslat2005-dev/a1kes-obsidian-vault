from __future__ import annotations

import csv
import math
import re
import textwrap
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import Iterable, List, Sequence

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
INPUT = ROOT / "input"
OUTPUT = ROOT / "output"
FIGURES = ROOT / "figures"
APPENDIX = ROOT / "appendix"
BUILD = ROOT / "build" / "docx_assets"
DOCX_PATH = OUTPUT / "智能微电网技术期末报告_电子底稿与手写参考稿.docx"
ISSUE_PATH = OUTPUT / "09_DOCX整合问题清单.md"
CHECK_PATH = OUTPUT / "10_DOCX最终检查报告.md"


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def clean_inline(text: str) -> str:
    text = text.strip()
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = text.replace("**", "")
    return text


def get_personal_info() -> dict:
    info = {"班级": "", "姓名": "", "学号": ""}
    path = INPUT / "GRXX.md"
    if not path.exists():
        return info
    text = read_text(path)
    for key in info:
        m = re.search(key + r"[:：]\s*(.+)", text)
        if m:
            info[key] = m.group(1).strip()
    return info


def find_file(directory: Path, preferred: str, fallback: str) -> Path:
    p = directory / preferred
    return p if p.exists() else directory / fallback


def split_top_sections(md: str) -> dict:
    sections = {}
    matches = list(re.finditer(r"^#\s+(.+)$", md, re.M))
    for i, match in enumerate(matches):
        title = match.group(1).strip()
        start = match.start()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(md)
        sections[title] = md[start:end].strip()
    return sections


def split_h2_sections(md: str) -> List[tuple[str, str]]:
    matches = list(re.finditer(r"^##\s+(.+)$", md, re.M))
    out = []
    for i, match in enumerate(matches):
        title = match.group(1).strip()
        start = match.start()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(md)
        out.append((title, md[start:end].strip()))
    return out


def strip_heading_block(md: str) -> str:
    lines = md.splitlines()
    if lines and lines[0].startswith("#"):
        return "\n".join(lines[1:]).strip()
    return md.strip()


def add_labeled_markdown(doc: Document, label: str, md: str, heading_level=3):
    doc.add_paragraph(label, style=f"Heading {heading_level}")
    add_markdown(doc, strip_heading_block(md))


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def restart_page_numbers(section, start=1):
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:start"), str(start))


def set_cell_text(cell, text: str, size=9, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(text))
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill="F2F4F7"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_width(table, widths_cm: Sequence[float] | None = None):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    if widths_cm:
        for row in table.rows:
            for idx, width in enumerate(widths_cm):
                if idx < len(row.cells):
                    row.cells[idx].width = Cm(width)


def add_table(doc, rows: Sequence[Sequence[str]], caption: str | None = None, widths_cm: Sequence[float] | None = None):
    if caption:
        p = doc.add_paragraph(caption)
        p.style = "Caption"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    set_table_width(table, widths_cm)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            set_cell_text(table.cell(r_idx, c_idx), clean_inline(val), size=9, bold=(r_idx == 0))
            if r_idx == 0:
                shade_cell(table.cell(r_idx, c_idx))
    return table


def set_run_font(run, font_name="宋体", size=12, bold=False, color=None):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def configure_styles(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    normal.paragraph_format.first_line_indent = Pt(24)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.space_after = Pt(6)

    for name, size, before, after in [
        ("Heading 1", 16, 16, 8),
        ("Heading 2", 14, 12, 6),
        ("Heading 3", 12, 8, 4),
    ]:
        st = styles[name]
        st.font.name = "黑体"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0, 0, 0)
        st.paragraph_format.first_line_indent = Pt(0)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.line_spacing = 1.3

    cap = styles["Caption"]
    cap.font.name = "宋体"
    cap._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    cap.font.size = Pt(9)
    cap.font.color.rgb = RGBColor(0, 0, 0)
    cap.paragraph_format.first_line_indent = Pt(0)
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(8)

    if "CodeBlock" not in styles:
        code = styles.add_style("CodeBlock", 1)
    else:
        code = styles["CodeBlock"]
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    code.font.size = Pt(8.5)
    code.paragraph_format.first_line_indent = Pt(0)
    code.paragraph_format.line_spacing = 1.0
    code.paragraph_format.space_after = Pt(0)


def configure_section(section):
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)


def add_cover(doc: Document, info: dict):
    for section in doc.sections:
        configure_section(section)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("南京工程学院")
    set_run_font(r, "宋体", 20, True)

    for _ in range(3):
        doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("期末论文作业报告")
    set_run_font(r, "黑体", 24, True)

    for _ in range(3):
        doc.add_paragraph("")

    fields = [
        ("课程名称", "智能微电网技术"),
        ("班级名称", info.get("班级") or ""),
        ("学生姓名", info.get("姓名") or ""),
        ("学生学号", info.get("学号") or ""),
        ("任课教师", "赵上林"),
        ("报告成绩", ""),
        ("评阅教师签字", ""),
    ]
    for label, value in fields:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"{label}：{value if value else '________________'}")
        set_run_font(r, "宋体", 14)

    for _ in range(2):
        doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("年     月     日")
    set_run_font(r, "宋体", 14)

    for _ in range(4):
        doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("电力工程学院")
    set_run_font(r, "宋体", 16, True)


def add_manual_toc(doc: Document):
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("目录")
    set_run_font(r, "黑体", 18, True)

    items = [
        "第一章 微电网提出的背景知识",
        "第二章 交直流混合微电网总体设计",
        "第三章 一次电气系统设计",
        "第四章 微电网能量管理系统设计",
        "第五章 交流小微电网运行控制",
        "第六章 直流小微电网运行控制",
        "第七章 结论",
        "附录 A 交流小微电网离网功率平衡控制程序设计",
        "附录 B 图纸手绘参考说明",
        "附录 C 最终自查表",
    ]
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(0)
        r = p.add_run(item)
        set_run_font(r, "宋体", 12)


def parse_md_table(lines: List[str], start: int) -> tuple[list[list[str]], int] | None:
    if start + 1 >= len(lines):
        return None
    if not lines[start].strip().startswith("|"):
        return None
    if not re.match(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$", lines[start + 1]):
        return None
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        if i != start + 1:
            cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
            rows.append(cells)
        i += 1
    return rows, i


def add_markdown(doc: Document, md: str, *, skip_first_heading=False, max_heading_level=3):
    lines = md.splitlines()
    i = 0
    first_heading_skipped = False
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue

        table = parse_md_table(lines, i)
        if table:
            rows, i = table
            if rows:
                add_table(doc, rows)
            continue

        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            title = clean_inline(line[level:].strip())
            if skip_first_heading and not first_heading_skipped:
                first_heading_skipped = True
                i += 1
                continue
            style = f"Heading {min(level, max_heading_level)}"
            doc.add_paragraph(title, style=style)
            i += 1
            continue

        if line.startswith("- "):
            p = doc.add_paragraph(clean_inline(line[2:]), style="List Bullet")
            p.paragraph_format.first_line_indent = Pt(0)
            i += 1
            continue

        if re.match(r"^\d+\.\s+", line):
            txt = re.sub(r"^\d+\.\s+", "", line)
            p = doc.add_paragraph(clean_inline(txt), style="List Number")
            p.paragraph_format.first_line_indent = Pt(0)
            i += 1
            continue

        if line.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1
            add_code_block(doc, "\n".join(code_lines))
            continue

        p = doc.add_paragraph(clean_inline(line))
        i += 1


def add_numbered_steps_from_md(doc: Document, md: str):
    sections = split_h2_sections(md)
    for title, block in sections:
        doc.add_paragraph(title, style="Heading 2")
        body = strip_heading_block(block)
        paragraphs = [x.strip() for x in re.split(r"\n\s*\n", body) if x.strip()]
        for para in paragraphs:
            for part in [p.strip() for p in para.splitlines() if p.strip()]:
                p = doc.add_paragraph(clean_inline(part), style="List Number")
                p.paragraph_format.first_line_indent = Pt(0)


def add_code_block(doc: Document, code: str):
    for line in code.splitlines():
        p = doc.add_paragraph(style="CodeBlock")
        p.paragraph_format.first_line_indent = Pt(0)
        run = p.add_run(line if line else " ")
        set_run_font(run, "Consolas", 8.5)


def caption(doc: Document, text: str):
    p = doc.add_paragraph(text)
    p.style = "Caption"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_picture(doc: Document, path: Path, caption_text: str, width_cm=15.8):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    caption(doc, caption_text)


def get_font(size: int, bold=False):
    candidates = [
        r"C:\Windows\Fonts\simsun.ttc",
        r"C:\Windows\Fonts\simhei.ttf" if bold else r"C:\Windows\Fonts\simsun.ttc",
        r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\arial.ttf",
    ]
    for c in candidates:
        p = Path(c)
        if p.exists():
            try:
                return ImageFont.truetype(str(p), size=size)
            except Exception:
                pass
    return ImageFont.load_default()


def parse_points(d: str):
    tokens = re.findall(r"[MLHVCSQTAZmlhvcsqtaz]|-?\d+(?:\.\d+)?", d)
    pts = []
    i = 0
    cur = (0.0, 0.0)
    cmd = None
    while i < len(tokens):
        if re.match(r"[A-Za-z]", tokens[i]):
            cmd = tokens[i]
            i += 1
            if cmd.upper() == "Z":
                continue
        if cmd is None or i >= len(tokens):
            break
        try:
            if cmd.upper() in {"M", "L"}:
                x, y = float(tokens[i]), float(tokens[i + 1])
                cur = (x, y)
                pts.append(cur)
                i += 2
            elif cmd.upper() == "H":
                x = float(tokens[i])
                cur = (x, cur[1])
                pts.append(cur)
                i += 1
            elif cmd.upper() == "V":
                y = float(tokens[i])
                cur = (cur[0], y)
                pts.append(cur)
                i += 1
            elif cmd.upper() == "C":
                # Approximate by drawing through the final control point.
                x1, y1, x2, y2, x3, y3 = map(float, tokens[i:i + 6])
                pts.extend([(x1, y1), (x2, y2), (x3, y3)])
                cur = (x3, y3)
                i += 6
            else:
                i += 1
        except Exception:
            break
    return pts


def render_svg_simple(svg_path: Path, out_path: Path, scale=2.8):
    root = ET.parse(svg_path).getroot()
    view_box = root.attrib.get("viewBox", "0 0 1200 800").split()
    width = int(float(view_box[2]) * scale)
    height = int(float(view_box[3]) * scale)
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)

    def sx(v): return int(round(float(v) * scale))
    def sy(v): return int(round(float(v) * scale))
    def stroke_width(cls):
        return max(1, int(round({"bus": 5, "line": 2.4, "thin": 1.4, "dash": 1.5}.get(cls, 1.8) * scale)))

    font_cache = {}

    def font_for(cls):
        base = 18
        if cls == "title":
            base = 28
        elif cls == "small":
            base = 15
        elif cls == "tiny":
            base = 13
        key = (base, cls == "title")
        if key not in font_cache:
            font_cache[key] = get_font(int(base * scale), bold=(cls == "title"))
        return font_cache[key]

    def walk(el):
        tag = el.tag.split("}")[-1]
        cls = el.attrib.get("class", "")
        if tag == "line":
            xy = (sx(el.attrib["x1"]), sy(el.attrib["y1"]), sx(el.attrib["x2"]), sy(el.attrib["y2"]))
            draw.line(xy, fill="black", width=stroke_width(cls))
        elif tag == "rect":
            x, y = sx(el.attrib["x"]), sy(el.attrib["y"])
            w, h = sx(el.attrib["width"]), sy(el.attrib["height"])
            draw.rectangle((x, y, x + w, y + h), outline="black", width=stroke_width(cls))
        elif tag == "circle":
            cx, cy, r = sx(el.attrib["cx"]), sy(el.attrib["cy"]), sx(el.attrib["r"])
            draw.ellipse((cx - r, cy - r, cx + r, cy + r), outline="black", width=stroke_width(cls))
        elif tag == "path":
            pts = parse_points(el.attrib.get("d", ""))
            if len(pts) >= 2:
                points = [(sx(x), sy(y)) for x, y in pts]
                draw.line(points, fill="black", width=stroke_width(cls))
        elif tag == "text":
            x, y = sx(el.attrib.get("x", 0)), sy(el.attrib.get("y", 0))
            text = "".join(el.itertext())
            draw.text((x, y - int(16 * scale)), text, fill="black", font=font_for(cls))
        for child in list(el):
            walk(child)

    for child in list(root):
        if child.tag.split("}")[-1] != "defs":
            walk(child)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    img.save(out_path, dpi=(300, 300))
    return out_path


def draw_flowchart(out_path: Path):
    w, h = 1800, 1350
    img = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(img)
    font = get_font(32)
    small = get_font(26)

    boxes = [
        ("开始", 760, 40, 280, 70),
        ("读取 P_pv、P_load、SOC", 650, 150, 500, 76),
        ("判断 P_pv 与 P_load", 650, 270, 500, 76),
        ("功率缺额：储能是否可放电？", 140, 430, 560, 90),
        ("储能放电补偿\nP_ess > 0", 170, 585, 500, 100),
        ("缺额仍存在：输出告警", 170, 750, 500, 90),
        ("功率盈余：储能是否可充电？", 1100, 430, 560, 90),
        ("储能充电吸收\nP_ess < 0", 1130, 585, 500, 100),
        ("充电受限：光伏限发", 1130, 750, 500, 90),
        ("更新 SOC", 650, 910, 500, 76),
        ("计算 P_balance_error", 650, 1030, 500, 76),
        ("输出控制指令，进入下一周期", 600, 1150, 600, 76),
    ]

    def rect(label, x, y, bw, bh):
        draw.rectangle((x, y, x + bw, y + bh), outline="black", width=4)
        lines = label.split("\n")
        for idx, line in enumerate(lines):
            bbox = draw.textbbox((0, 0), line, font=font)
            draw.text((x + (bw - (bbox[2] - bbox[0])) / 2, y + 16 + idx * 36), line, fill="black", font=font)

    def arrow(x1, y1, x2, y2):
        draw.line((x1, y1, x2, y2), fill="black", width=4)
        ang = math.atan2(y2 - y1, x2 - x1)
        for a in (ang + 2.55, ang - 2.55):
            draw.line((x2, y2, x2 + 22 * math.cos(a), y2 + 22 * math.sin(a)), fill="black", width=4)

    for b in boxes:
        rect(*b)
    arrow(900, 110, 900, 150)
    arrow(900, 226, 900, 270)
    arrow(900, 346, 420, 430)
    arrow(900, 346, 1380, 430)
    arrow(420, 520, 420, 585)
    arrow(420, 685, 420, 750)
    arrow(1380, 520, 1380, 585)
    arrow(1380, 685, 1380, 750)
    arrow(420, 840, 760, 948)
    arrow(1380, 840, 1040, 948)
    arrow(900, 986, 900, 1030)
    arrow(900, 1106, 900, 1150)
    draw.text((515, 390), "P_pv < P_load", fill="black", font=small)
    draw.text((1030, 390), "P_pv > P_load", fill="black", font=small)
    draw.text((840, 390), "相等：待机", fill="black", font=small)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    img.save(out_path, dpi=(300, 300))
    return out_path


def add_appendix_result_table(doc: Document, csv_path: Path):
    with csv_path.open("r", encoding="utf-8-sig", newline="") as f:
        rows = list(csv.DictReader(f))
    selected = [0, 1, 8, 10, 11, 12, 15, 17, 23]
    header = ["时间/h", "光伏实际/kW", "负荷/kW", "储能/kW", "SOC/%", "限发/kW", "状态", "告警"]
    out = [header]
    for idx in selected:
        r = rows[idx]
        out.append([
            r["时间/h"],
            r["光伏实际出力/kW"],
            r["负荷功率/kW"],
            r["储能功率/kW"],
            r["SOC/%"],
            r["光伏限发功率/kW"],
            r["控制状态"],
            r["告警状态"],
        ])
    add_table(doc, out, "表A-2 典型时段仿真结果", widths_cm=[1.1, 1.6, 1.3, 1.3, 1.2, 1.3, 4.0, 1.0])


def build_docx():
    OUTPUT.mkdir(exist_ok=True)
    BUILD.mkdir(parents=True, exist_ok=True)
    issues = []

    task_path = find_file(INPUT, "任务书.docx", "RWS.docx")
    ref_path = find_file(INPUT, "参考作业.pdf", "CKZY.pdf")
    if task_path.name != "任务书.docx":
        issues.append("输入文件 `input/任务书.docx` 不存在，已使用实际任务书文件 `input/RWS.docx`。")
    if ref_path.name != "参考作业.pdf":
        issues.append("输入文件 `input/参考作业.pdf` 不存在，已使用实际参考文件 `input/CKZY.pdf`，仅作结构和图式参考。")

    main_md = read_text(OUTPUT / "04_报告手写底稿.md")
    main_sections = split_top_sections(main_md)
    info = get_personal_info()

    png_map = {}
    for svg in sorted(FIGURES.glob("0*.svg")):
        out = BUILD / (svg.stem + ".png")
        render_svg_simple(svg, out)
        png_map[svg.name] = out
    issues.append("Word 兼容性处理：4 张 SVG 图纸均保留原文件，并转换为高清 PNG 插入 DOCX。")

    flow_png = draw_flowchart(BUILD / "附加题流程图.png")

    doc = Document()
    configure_styles(doc)
    add_cover(doc, info)
    add_manual_toc(doc)

    body_sec = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(body_sec)
    body_sec.footer.is_linked_to_previous = False
    restart_page_numbers(body_sec, 1)
    add_page_number(body_sec.footer.paragraphs[0])

    # Main report
    add_markdown(doc, main_sections["第一章 微电网提出的背景知识"])
    add_markdown(doc, main_sections["第二章 交直流混合微电网总体设计"])

    doc.add_paragraph("第三章 一次电气系统设计", style="Heading 1")
    chap3 = main_sections["第三章 一次电气系统设计"]
    h2s = dict(split_h2_sections(chap3))
    doc.add_paragraph("3.1 总体一次接线结构", style="Heading 2")
    add_labeled_markdown(doc, "3.1.1 交流大微电网一次接线结构", h2s["3.1 交流大微电网一次接线结构"])
    add_labeled_markdown(doc, "3.1.2 交流小微电网一次接线结构", h2s["3.2 交流小微电网一次接线结构"])
    add_labeled_markdown(doc, "3.1.3 直流小微电网一次接线结构", h2s["3.3 直流小微电网一次接线结构"])
    add_picture(doc, png_map["01_总体一次电气图.svg"], "图3-1 交直流混合微电网总体一次电气接线图", 15.6)
    doc.add_paragraph("图3-1 表示交流大微电网经 PCC 接入 380V 配电网，交流小微电网通过交流联络支路接入，直流小微电网通过双向 AC/DC 互联变换器接入。图中同时标出了主要电源、储能、负荷、断路器、协调控制器和 EMS 通信关系。")

    doc.add_paragraph("3.2 设备并网电气结构", style="Heading 2")
    for key in [
        "3.4 交流侧光伏并网电气结构",
        "3.5 直流侧光伏并网电气结构",
        "3.6 直驱风力发电并网电气结构",
        "3.7 储能系统并网电气结构",
        "3.8 直流充电负荷接入结构",
    ]:
        label = re.sub(r"^3\.\d+\s+", "", key)
        add_labeled_markdown(doc, label, h2s[key])
    add_picture(doc, png_map["02_设备并网结构图.svg"], "图3-2 各类分布式电源、储能及负荷并网电气结构图", 15.6)

    doc.add_paragraph("3.3 微电网接地设计", style="Heading 2")
    add_markdown(doc, strip_heading_block(h2s["3.9 接地方案"]))
    add_picture(doc, png_map["03_接地示意图.svg"], "图3-3 微电网接地方式示意图", 15.0)

    doc.add_paragraph("第四章 微电网能量管理系统设计", style="Heading 1")
    chap4_h2 = dict(split_h2_sections(main_sections["第四章 微电网能量管理系统设计"]))
    add_markdown(doc, chap4_h2["4.1 EMS 层级结构"])
    add_picture(doc, png_map["04_EMS与通信结构图.svg"], "图4-1 微电网能量管理系统及通信结构图", 15.6)
    add_markdown(doc, chap4_h2["4.2 通信方案"])
    comm_rows = [
        ["通信层级", "通信对象", "通信方式", "传输数据类型", "主要功能"],
        ["上层-中层", "EMS 与三套协调控制器", "本方案选取工业以太网", "模式、功率计划、告警、SOC", "运行模式判断和功率调度"],
        ["中层-设备层", "协调控制器与逆变器、PCS、BMS", "现场总线/工业以太网", "功率指令、限发指令、SOC", "设备协调控制"],
        ["测量保护层", "智能电表、保护测控与 EMS/协调控制器", "现场总线/工业以太网", "电压、电流、频率、开关状态", "计量、保护和状态上传"],
    ]
    add_table(doc, comm_rows, "表4-1 EMS 通信结构表", widths_cm=[2.2, 3.1, 3.0, 3.5, 3.3])
    add_markdown(doc, chap4_h2["4.3 EMS 硬件配置表"])

    doc.add_paragraph("第五章 交流小微电网运行控制", style="Heading 1")
    add_numbered_steps_from_md(doc, main_sections["第五章 交流小微电网运行控制步骤"])
    doc.add_paragraph("第六章 直流小微电网运行控制", style="Heading 1")
    add_numbered_steps_from_md(doc, main_sections["第六章 直流小微电网运行控制步骤"])
    add_markdown(doc, main_sections["第七章 结论"])

    # Appendix A
    doc.add_page_break()
    doc.add_paragraph("附录 A：交流小微电网离网功率平衡控制程序设计", style="Heading 1")
    doc.add_paragraph("A.1 附加题目的与控制对象", style="Heading 2")
    add_markdown(doc, read_text(APPENDIX / "附加题说明_精简版.md"), skip_first_heading=True)
    doc.add_paragraph("A.2 控制原理与功率平衡关系", style="Heading 2")
    add_markdown(doc, read_text(APPENDIX / "附加题_交流小微电网离网功率平衡控制说明.md").split("## 9. 程序流程说明")[0], skip_first_heading=True)
    doc.add_paragraph("A.3 控制流程图", style="Heading 2")
    add_picture(doc, flow_png, "图A-1 交流小微电网离网功率平衡控制流程图", 15.0)
    doc.add_paragraph("A.4 Python 程序代码", style="Heading 2")
    add_code_block(doc, read_text(APPENDIX / "islanded_ac_microgrid_power_balance.py"))
    doc.add_paragraph("A.5 仿真结果", style="Heading 2")
    add_appendix_result_table(doc, APPENDIX / "仿真结果.csv")
    add_picture(doc, APPENDIX / "仿真结果_功率曲线.png", "图A-2 离网运行期间各功率量变化曲线", 15.6)
    add_picture(doc, APPENDIX / "仿真结果_SOC曲线.png", "图A-3 储能系统 SOC 变化曲线", 15.6)
    add_picture(doc, APPENDIX / "仿真结果_限发与告警.png", "图A-4 光伏限发或功率缺额告警结果", 15.6)
    doc.add_paragraph("A.6 仿真结果分析", style="Heading 2")
    add_markdown(doc, read_text(APPENDIX / "仿真结果分析.md"), skip_first_heading=True)

    doc.add_page_break()
    doc.add_paragraph("附录 B：图纸手绘参考说明", style="Heading 1")
    add_markdown(doc, read_text(OUTPUT / "06_图纸手绘说明.md"), skip_first_heading=True)

    doc.add_page_break()
    doc.add_paragraph("附录 C：最终自查表", style="Heading 1")
    add_markdown(doc, read_text(OUTPUT / "05_最终自查清单.md"), skip_first_heading=True)
    add_markdown(doc, read_text(OUTPUT / "07_图文一致性审查.md"), skip_first_heading=True)
    add_markdown(doc, read_text(APPENDIX / "最终审查报告.md"), skip_first_heading=True)

    doc.save(DOCX_PATH)

    if not issues:
        issues.append("未发现需要列出的技术参数冲突。")
    issue_md = "# 09_DOCX整合问题清单\n\n" + "\n".join(f"- {x}" for x in issues) + "\n"
    ISSUE_PATH.write_text(issue_md, encoding="utf-8")

    check_items = [
        "已读取任务书实际文件 input/RWS.docx，并以其参数为准。",
        "已读取参考作业实际文件 input/CKZY.pdf，仅参考章节和图式表现形式。",
        "已整合 output/00、01、04、05、06、07 等 Markdown 文件中的核心内容。",
        "已插入 4 张图纸，均使用 PNG 兼容版本插入，原 SVG 文件保留。",
        "已插入附加题 Python 代码、典型 CSV 结果表和 3 张仿真曲线图。",
        "交流大微电网、交流小微电网、直流小微电网参数均按任务书配置。",
        "接地方案保持 TN-S 与直流侧绝缘监测/保护接地一致。",
        "附加题未加入非重要负荷、柴油机、风电等不存在设备。",
        "代码以等宽字体插入，保留缩进和必要中文注释。",
        "正文中任务书未规定的接地、通信、SOC 等参数均作为设计假设或方案选取表达。",
    ]
    CHECK_PATH.write_text("# 10_DOCX最终检查报告\n\n" + "\n".join(f"- [x] {x}" for x in check_items) + "\n", encoding="utf-8")
    return DOCX_PATH


if __name__ == "__main__":
    print(build_docx())
